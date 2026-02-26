from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('StoreAI Enterprise: Role-Based Workflows', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('This document outlines the key workflows, architecture, and technology stack for the StoreAI Enterprise system.')
    
    # Tech Stack
    doc.add_heading('1. Tech Stack & Architecture', level=1)
    
    doc.add_heading('Technology Stack', level=2)
    p = doc.add_paragraph()
    p.add_run('Frontend: ').bold = True
    p.add_run('React 18, Vite, TypeScript, Tailwind CSS, Lucide React.')
    
    p = doc.add_paragraph()
    p.add_run('Backend: ').bold = True
    p.add_run('Node.js, Express, TypeScript, Prisma ORM.')
    
    p = doc.add_paragraph()
    p.add_run('Database: ').bold = True
    p.add_run('PostgreSQL (Relational).')
    
    doc.add_heading('System Architecture', level=2)
    doc.add_paragraph('The system follows a Monolithic Layered Architecture with clear separation of concerns: Routes -> Middleware -> Controllers -> Services -> Data Access (Prisma). It supports Multi-tenancy via database-level isolation.')

    doc.add_page_break()
    
    # Workflows
    doc.add_heading('2. Role-Based Workflows', level=1)
    
    # HR - Staffing
    add_workflow_section(doc, 
        "HR Manager - Staffing (Add Employee)",
        "Navigate to Workforce > Employee Master. Click 'Add Employee' to register new staff details.",
        "hr_employee_master.png")

    # HR - Attendance
    add_workflow_section(doc, 
        "HR Manager - Attendance & Performance",
        "Navigate to Attendance Master. Mark daily presence and assign 1-5 star performance ratings.",
        "hr_attendance_performance.png")

    # Procurement
    add_workflow_section(doc, 
        "Procurement Manager - Purchase Order",
        "Navigate to Procurement Hub > New Order. Select Supplier and Products to generate a PO.",
        "procurement_purchase_order.png")

    # Warehouse
    add_workflow_section(doc, 
        "Warehouse Manager - Inbound (GRN)",
        "Receive goods against an approved PO. Confirm quantity and batch numbers.",
        "warehouse_inbound_grn.png")

    # Sales
    add_workflow_section(doc, 
        "Sales Staff - Home Delivery (POS)",
        "Process a sale at the POS Terminal. Select 'Home Delivery' option to capture address details.",
        "sales_pos_terminal.png")

    # Logistics
    add_workflow_section(doc, 
        "Logistics Coordinator - Outbound Dispatch",
        "Track Home Delivery orders. Enter Tracking Number and Carrier details for dispatch.",
        "logistics_outbound_dispatch.png")

    # Admin
    add_workflow_section(doc, 
        "System Administrator - Access Control",
        "Manage Users, Roles, and Permissions in the Tenant Settings dashboard.",
        "admin_system_access.png")
        
    output_path = "StoreAI_Role_Workflows.docx"
    doc.save(output_path)
    print(f"Document saved: {os.path.abspath(output_path)}")

def add_workflow_section(doc, title, description, image_filename):
    doc.add_heading(title, level=2)
    doc.add_paragraph(description)
    
    if os.path.exists(image_filename):
        try:
            doc.add_picture(image_filename, width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Image could not be loaded: {e}]")
            
    doc.add_paragraph("") # Spacing

if __name__ == "__main__":
    create_document()
